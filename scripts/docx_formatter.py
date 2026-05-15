"""
Rich DOCX extraction module for wiki integrator.

Converts .docx files to Markdown while preserving:
- Heading levels (Heading 1 -> h1, Heading 2 -> h2, etc.)
- Bullet and numbered lists
- Bold, italic, underline text styling
- Tables (as Markdown tables)
- Inline code formatting
- Paragraph spacing and structure
"""

import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


# Mapping of Word styles to Markdown heading levels
HEADING_STYLE_MAP = {
    'Heading 1': 1,
    'Heading 2': 2,
    'Heading 3': 3,
    'Heading 4': 4,
    'Heading 5': 5,
    'Heading 6': 6,
    'Heading 7': 7,
    'Heading 8': 8,
    'Heading 9': 9,
}

# Styles that should be treated as list items
BULLET_STYLE_NAMES = {'List Bullet', 'List Bullet 2', 'List Bullet 3'}
NUMBERED_STYLE_NAMES = {'List Number', 'List Number 2', 'List Number 3'}
INDENTED_STYLE_NAMES = {'List Continue', 'List Continue 2', 'List Continue 3'}

# Style that removes spacing
NO_SPACING_STYLE = 'No Spacing'


def _get_indent_level(para):
    """Get the indent level of a paragraph (for list nesting)."""
    pr = para.paragraph_format
    left_indent = pr.left_indent
    if left_indent is None:
        return 0
    return int((left_indent / 0.01) // 400)  # Word uses 20ths of a point


def _get_paragraph_text(para):
    """Extract text from a paragraph, preserving inline formatting markers."""
    parts = []
    for run in para.runs:
        text = run.text
        if not text:
            continue
        
        # Check inline formatting
        bold = run.bold
        italic = run.italic
        underline = run.underline
        
        # Apply formatting markers
        if bold and italic:
            text = f'***{text}***'
        elif bold:
            text = f'**{text}**'
        elif italic:
            text = f'*{text}*'
        elif underline:
            text = f'__{text}__'
        
        parts.append(text)
    
    return ''.join(parts)


def _get_paragraph_style_name(para):
    """Get the style name of a paragraph."""
    if para.style is None:
        return 'Normal'
    return para.style.name


def _paragraph_to_heading(para):
    """Convert a heading paragraph to Markdown heading."""
    style_name = _get_paragraph_style_name(para)
    level = HEADING_STYLE_MAP.get(style_name)
    if level is None:
        return None
    
    text = _get_paragraph_text(para)
    prefix = '#' * level
    return f'{prefix} {text}\n'


def _paragraph_to_bullet(para):
    """Convert a bullet list paragraph to Markdown."""
    style_name = _get_paragraph_style_name(para)
    indent = _get_indent_level(para)
    
    # Determine bullet character based on indent level
    bullets = ['-', '+', '*']
    bullet_char = bullets[min(indent, len(bullets) - 1)]
    prefix = ' ' * (indent * 2) + bullet_char + ' '
    
    text = _get_paragraph_text(para)
    return f'{prefix}{text}\n'


def _paragraph_to_numbered(para, number=1):
    """Convert a numbered list paragraph to Markdown."""
    style_name = _get_paragraph_style_name(para)
    indent = _get_indent_level(para)
    
    prefix = ' ' * (indent * 2) + f'{number}. '
    text = _get_paragraph_text(para)
    return f'{prefix}{text}\n'


def _paragraph_to_normal(para):
    """Convert a normal paragraph to Markdown text."""
    text = _get_paragraph_text(para)
    
    # Handle alignment
    alignment = para.alignment
    if alignment == WD_ALIGN_PARAGRAPH.CENTER:
        text = f'\n{text}\n'
    elif alignment == WD_ALIGN_PARAGRAPH.RIGHT:
        text = f'\n{text}\n'
    
    return f'{text}\n'


def _paragraph_to_table_row(cell):
    """Convert a table cell to Markdown table content."""
    # Get all paragraphs in the cell (some cells have multi-line content)
    parts = []
    for para in cell.paragraphs:
        text = _get_paragraph_text(para).strip()
        if text:
            parts.append(text)
    return '\n'.join(parts)


def _table_to_markdown(table):
    """Convert a docx table to Markdown format."""
    if not table.rows:
        return ''
    
    rows = list(table.rows)
    markdown_lines = []
    
    # Calculate column widths for alignment (optional, simplified)
    num_cols = len(rows[0].cells)
    
    for i, row in enumerate(rows):
        cells = []
        for cell in row.cells:
            content = _paragraph_to_table_row(cell)
            # Escape pipe characters within cells
            content = content.replace('|', '\\|')
            cells.append(content.strip())
        
        # Limit cells to match first row column count
        cells = cells[:num_cols]
        
        if i == 0:
            # Header row (use bold by default for first row)
            markdown_lines.append('| ' + ' | '.join(f'**{c}**' for c in cells) + ' |')
            markdown_lines.append('| ' + ' | '.join(['---'] * len(cells)) + ' |')
        else:
            markdown_lines.append('| ' + ' | '.join(cells) + ' |')
    
    return '\n' + '\n'.join(markdown_lines) + '\n\n'


def docx_to_markdown(docx_path):
    """
    Convert a .docx file to rich Markdown.
    
    Args:
        docx_path: Path to the .docx file.
        
    Returns:
        A string containing the Markdown representation with preserved
        headings, lists, formatting, and tables.
    """
    doc = Document(docx_path)
    result_parts = []
    
    # Track list state
    current_list_type = None  # 'bullet', 'numbered', or None
    list_counter = 0
    
    for para in doc.paragraphs:
        style_name = _get_paragraph_style_name(para)
        text_content = para.text.strip()
        
        # Skip empty paragraphs that have no styling significance
        if not text_content and style_name not in (NO_SPACING_STYLE,):
            # Check if this paragraph has formatting that matters
            has_formatting = any(
                r.bold or r.italic or r.underline
                for r in para.runs
            )
            if has_formatting:
                # Keep formatting but note it's empty content
                formatted = _get_paragraph_text(para)
                result_parts.append(f'{formatted}\n')
            continue
        
        # Heading styles
        if style_name in HEADING_STYLE_MAP:
            # Close any open lists
            current_list_type = None
            list_counter = 0
            result_parts.append(_paragraph_to_heading(para))
            result_parts.append('\n')
            continue
        
        # Bullet list items
        if style_name in BULLET_STYLE_NAMES:
            current_list_type = 'bullet'
            result_parts.append(_paragraph_to_bullet(para))
            continue
        
        # Numbered list items
        if style_name in NUMBERED_STYLE_NAMES:
            if current_list_type != 'numbered':
                current_list_type = 'numbered'
                list_counter = 0
            list_counter += 1
            result_parts.append(_paragraph_to_numbered(para, list_counter))
            continue
        
        # List continue styles (indented items in lists)
        if style_name in INDENTED_STYLE_NAMES:
            if current_list_type == 'bullet':
                result_parts.append(_paragraph_to_bullet(para))
            elif current_list_type == 'numbered':
                list_counter += 1
                result_parts.append(_paragraph_to_numbered(para, list_counter))
            continue
        
        # No spacing style - keep content but don't add extra blank lines
        if style_name == NO_SPACING_STYLE:
            result_parts.append(_get_paragraph_text(para))
            result_parts.append('\n')
            continue
        
        # Normal paragraph
        current_list_type = None
        list_counter = 0
        result_parts.append(_paragraph_to_normal(para))
    
    # Process tables separately (they need full table context)
    for table in doc.tables:
        table_md = _table_to_markdown(table)
        result_parts.append(table_md)
    
    return ''.join(result_parts)


def docx_to_markdown_simplified(docx_path):
    """
    Simplified DOCX-to-Markdown conversion.
    
    Falls back to basic text extraction when rich formatting detection fails.
    Use this as a fallback when docx_to_markdown raises an exception.
    """
    doc = Document(docx_path)
    
    # Use the rich formatter, but fall back to paragraph text if needed
    try:
        return docx_to_markdown(docx_path)
    except Exception:
        # Fallback: just get paragraph text
        return '\n\n'.join(p.text for p in doc.paragraphs if p.text.strip())