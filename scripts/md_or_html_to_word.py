#!/usr/bin/env python3
"""
Convert Markdown (.md) or HTML (.html) wiki pages to Word (.docx) format.

Designed for the Living LLM Wiki workflow. Supports:
- Automatic wiki/ prefix resolution for convenience
- YAML frontmatter stripping
- Full Markdown element support (tables, code blocks, links, etc.)
- Batch conversion of all wiki pages
- Configurable output directory

Usage:
    # Convert a single file (full path)
    python scripts/md_or_html_to_word.py wiki/log.md export/log.docx

    # Convert by name (auto-resolves to wiki/log.md)
    python scripts/md_or_html_to_word.py log export/log.docx

    # Convert with defaults (output goes to export/log.docx)
    python scripts/md_or_html_to_word.py log

    # Batch convert all wiki pages
    python scripts/md_or_html_to_word.py --batch

    # Batch convert from specific directory
    python scripts/md_or_html_to_word.py --batch --source wiki --output export
"""

import sys
import os
import argparse
import re
import shutil

# --- Dependency Check Layer ---
missing_modules = []

try:
    import markdown  # main conversion engine
except ImportError:
    missing_modules.append("markdown")

try:
    from bs4 import BeautifulSoup
except ImportError:
    missing_modules.append("beautifulsoup4")

try:
    from docx import Document
except ImportError:
    missing_modules.append("python-docx")

try:
    import lxml  # parser used by BeautifulSoup
except ImportError:
    missing_modules.append("lxml")


# If anything missing → notify user clearly and exit
if missing_modules:
    print("\n[ERROR] Missing required Python modules:\n")
    for module in missing_modules:
        print(f"   - {module}")
    print("\n[INFO] Install them using:\n")
    print(f"   pip install {' '.join(missing_modules)}\n")
    print("[INFO] Tip: If pip doesn't work, try:")
    print("   python -m pip install " + " ".join(missing_modules))
    print()
    sys.exit(1)

# --- Imports (safe after check) ---
import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# --- Configuration ---
DEFAULT_OUTPUT_DIR = "export"
DEFAULT_SOURCE_DIR = "wiki"
SUPPORTED_MD_EXTENSIONS = [
    "tables",
    "fenced_code",
    "codehilite",
    "attr_list",
    "md_in_html",
    "toc",
    "sane_lists",
]


# --- Frontmatter Handling ---
def strip_frontmatter(content):
    """Remove YAML frontmatter from the top of a file if present."""
    pattern = r"^---\n.*?^---\n*"
    result = re.sub(pattern, "", content, count=1, flags=re.MULTILINE | re.DOTALL)
    # Clean up any leading/trailing blank lines left by frontmatter removal
    result = re.sub(r"^\n+", "", result)
    return result


# --- Markdown to HTML Preprocessing ---
def preprocess_markdown_to_html(md_content):
    """Convert Markdown to HTML with full extension support."""
    # Strip frontmatter first
    md_content = strip_frontmatter(md_content)

    return markdown.markdown(
        md_content,
        extensions=SUPPORTED_MD_EXTENSIONS,
        extension_configs={
            "codehilite": {
                "guess_lang": False,
                "css_class": "highlight",
            }
        },
    )


# --- HTML Cleanup ---
def clean_html_for_conversion(html_content):
    """Clean and normalize HTML for docx conversion."""
    soup = BeautifulSoup(html_content, "lxml")

    # Remove script and style elements
    for element in soup.find_all(["script", "style", "meta", "link"]):
        element.decompose()

    # Fix SilverBullet-style links [[page_name]] in anchor tags
    for a_tag in soup.find_all("a"):
        href = a_tag.get("href", "")
        if href and href.startswith("[["):
            # SilverBullet live query or link syntax
            display_text = a_tag.get_text().strip()
            if display_text:
                a_tag.string = display_text
            else:
                a_tag.decompose()

    # Remove class attributes that cause docx issues
    for element in soup.find_all(True):
        if "class" in element.attrs:
            del element["class"]

    return str(soup)


# --- Formatting Helpers ---
def add_formatted_text(paragraph, element):
    """Add text content with inline formatting (bold, italic, underline, etc.)."""
    if element is None:
        return

    for content in element.children:
        if content.name is None:
            # Text node
            text = content.get_text()
            if text.strip():
                paragraph.add_run(text)
        elif isinstance(content, str):
            text = content.strip()
            if text:
                paragraph.add_run(text)
        else:
            tag = content.name.lower()
            text = content.get_text()
            if not text.strip():
                continue

            run = paragraph.add_run(text)

            # Apply formatting based on tag
            if tag in ["b", "strong"]:
                run.bold = True
            if tag in ["i", "em"]:
                run.italic = True
            if tag == "u":
                run.underline = True
            if tag == "s" or tag == "strike" or tag == "del":
                run.strike = True
            if tag == "code":
                run.font.name = "Consolas"
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x83, 0x1C, 0x1C)  # Dark red


def process_heading(text, soup, level, doc):
    """Process headings with proper level mapping and cleanup."""
    doc.add_heading(text, level=level)


def process_code_block(element, doc):
    """Process <pre><code> blocks with proper formatting."""
    code = element.find("code")
    if code is None:
        code = element
    code_text = code.get_text()

    # Add as a paragraph with monospace font
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)


def process_table(table_element, doc):
    """Convert an HTML table to a docx table with proper styling."""
    rows = table_element.find_all("tr")
    if not rows:
        return

    cols = rows[0].find_all(["td", "th"])
    if not cols:
        return

    # Determine if first row is a header
    first_row_headers = cols[0].find_all("th")
    has_header = len(first_row_headers) > 0

    num_rows = len(rows)
    num_cols = len(cols)

    table = doc.add_table(rows=num_rows, cols=num_cols, style="Table Grid")

    for i, row in enumerate(rows):
        cells = row.find_all(["td", "th"])
        is_header_row = (i == 0 and has_header)

        for j, cell in enumerate(cells):
            cell_text = cell.get_text().strip()
            cell_elem = table.rows[i].cells[j]

            # Clear default paragraph
            cell_elem.paragraphs[0].clear()

            if is_header_row:
                run = cell_elem.paragraphs[0].add_run(cell_text)
                run.bold = True
                run.font.size = Pt(11)
            else:
                add_formatted_text(cell_elem.paragraphs[0], cell)
                cell_elem.paragraphs[0].paragraph_format.space_before = Pt(2)
                cell_elem.paragraphs[0].paragraph_format.space_after = Pt(2)

    # (Table spacing is handled per-cell above)


def process_blockquote(blockquote, doc):
    """Process HTML blockquote as indented paragraph."""
    text = blockquote.get_text().strip()
    if text:
        p = doc.add_paragraph(text)
        p.paragraph_format.left_indent = Pt(36)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        # Add cite reference if present
        cite = blockquote.find("cite")
        if cite:
            p2 = doc.add_paragraph(f"— {cite.get_text()}")
            p2.paragraph_format.left_indent = Pt(36)
            p2.paragraph_format.space_after = Pt(8)


# --- Core HTML to DOCX Conversion ---
def html_to_docx(html_content, doc):
    """Convert HTML content to docx Document object."""
    soup = BeautifulSoup(html_content, "lxml")
    body = soup.body if soup.body else soup

    for element in body.children:
        if element.name is None:
            # Raw text
            text = element.strip()
            if text:
                doc.add_paragraph(text)
            continue

        tag = element.name.lower()

        # Headings
        if tag in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            level = int(tag[1])
            heading_text = element.get_text().strip()
            doc.add_heading(heading_text, level=min(level, 6))

        # Paragraphs
        elif tag == "p":
            add_formatted_text(doc.add_paragraph(), element)

        # Horizontal rule
        elif tag == "hr":
            doc.add_page_break()

        # Lists
        elif tag == "ul":
            for li in element.find_all("li", recursive=False):
                text = li.get_text().strip()
                if text:
                    doc.add_paragraph(text, style="List Bullet")

        elif tag == "ol":
            for i, li in enumerate(element.find_all("li", recursive=False), 1):
                text = li.get_text().strip()
                if text:
                    doc.add_paragraph(text, style="List Number")

        # Code blocks
        elif tag == "pre":
            process_code_block(element, doc)

        # Tables
        elif tag == "table":
            process_table(element, doc)

        # Blockquotes
        elif tag == "blockquote":
            process_blockquote(element, doc)

        # Images
        elif tag == "img":
            src = element.get("src", "")
            alt = element.get("alt", "")
            if alt:
                p = doc.add_paragraph(f"[Image: {alt}]")
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
            elif src:
                p = doc.add_paragraph(f"[Image: {src}]")
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)

        # Div blocks (treat children as paragraphs)
        elif tag == "div":
            for child in element.children:
                if child.name and child.name.lower() in ["p", "h1", "h2", "h3", "h4", "h5", "h6", "ul", "ol", "pre", "table", "blockquote", "hr"]:
                    child_copy = BeautifulSoup(str(child), "lxml").find()
                    # Process known block elements by re-invoking logic
                    child_tag = child.name.lower()
                    if child_tag == "p":
                        add_formatted_text(doc.add_paragraph(), child)
                    elif child_tag in ["h1", "h2", "h3", "h4", "h5", "h6"]:
                        level = int(child_tag[1])
                        doc.add_heading(child.get_text().strip(), level=min(level, 6))
                    elif child_tag == "pre":
                        process_code_block(child, doc)
                    elif child_tag == "table":
                        process_table(child, doc)
                    elif child_tag == "blockquote":
                        process_blockquote(child, doc)
                    elif child_tag == "hr":
                        doc.add_page_break()
                    elif child_tag == "ul":
                        for li in child.find_all("li", recursive=False):
                            text = li.get_text().strip()
                            if text:
                                doc.add_paragraph(text, style="List Bullet")
                    elif child_tag == "ol":
                        for i, li in enumerate(child.find_all("li", recursive=False), 1):
                            text = li.get_text().strip()
                            if text:
                                doc.add_paragraph(text, style="List Number")

        # Skip other block elements we don't specially handle
        # (they'll be processed by BeautifulSoup's recursive content)


# --- Input Loader ---
def resolve_input_path(input_path):
    """Resolve a possibly ambiguous input path to an absolute file path.

    Priority:
    1. Absolute path
    2. Relative path as-is
    3. wiki/<name> directory resolution
    """
    # Already absolute
    if os.path.isabs(input_path):
        if os.path.isfile(input_path):
            return input_path
        return None

    # Direct file exists
    if os.path.isfile(input_path):
        return input_path

    # Try wiki/ prefix
    wiki_path = os.path.join(DEFAULT_SOURCE_DIR, input_path)
    if os.path.isfile(wiki_path):
        return wiki_path

    # Try adding common extensions to wiki/
    for ext in [".md", ".markdown", ".html", ".htm"]:
        wiki_ext_path = wiki_path + ext
        if os.path.isfile(wiki_ext_path):
            return wiki_ext_path

    return None


def load_input_file(input_path):
    """Load and prepare file content for conversion. Returns HTML string."""
    resolved_path = resolve_input_path(input_path)

    if resolved_path is None:
        raise FileNotFoundError(
            f"File not found: {input_path}\n"
            f"  Searched: {input_path}, wiki/{input_path}, wiki/{input_path}.md, wiki/{input_path}.html"
        )

    ext = os.path.splitext(resolved_path)[1].lower()

    try:
        with open(resolved_path, "r", encoding="utf-8") as f:
            content = f.read()
    except Exception as e:
        raise RuntimeError(f"Failed to read file '{resolved_path}': {e}")

    if ext in [".md", ".markdown"]:
        return preprocess_markdown_to_html(content)
    elif ext in [".html", ".htm"]:
        return content
    else:
        raise ValueError(
            f"Unsupported file type '.{ext[1:]}'. Use .md, .markdown, .html, or .htm"
        )


# --- Output ---
def ensure_output_dir(output_path):
    """Ensure the output directory exists."""
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir, exist_ok=True)
    print(f"   [CREATED] {output_dir}/")


def convert_to_docx(input_path, output_path):
    """Convert a single file from markdown/html to docx."""
    print(f"\n{'='*50}")
    print(f"  Converting: {input_path}")
    print(f"  Output:     {output_path}")
    print(f"{'='*50}")

    # Load and convert
    html_content = load_input_file(input_path)

    # Clean HTML for docx compatibility
    html_content = clean_html_for_conversion(html_content)

    # Create docx document
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Convert HTML to docx content
    html_to_docx(html_content, doc)

    # Ensure output directory exists
    ensure_output_dir(output_path)

    # Save
    doc.save(output_path)
    print(f"  [SUCCESS] Saved: {output_path}")
    return True


# --- Batch Conversion ---
def batch_convert(source_dir=None, output_dir=None):
    """Convert all wiki pages in source_dir to docx in output_dir."""
    source_dir = source_dir or DEFAULT_SOURCE_DIR
    output_dir = output_dir or DEFAULT_OUTPUT_DIR

    if not os.path.isdir(source_dir):
        print(f"[ERROR] Source directory not found: {source_dir}")
        sys.exit(1)

    # Collect files
    md_files = []
    html_files = []

    for entry in sorted(os.listdir(source_dir)):
        full_path = os.path.join(source_dir, entry)
        if not os.path.isfile(full_path):
            continue
        ext = os.path.splitext(entry)[1].lower()
        if ext in [".md", ".markdown"]:
            md_files.append(entry)
        elif ext in [".html", ".htm"]:
            html_files.append(entry)

    total = len(md_files) + len(html_files)
    if total == 0:
        print(f"\nNo markdown or HTML files found in {source_dir}/")
        return

    print(f"\n{'='*50}")
    print(f"  BATCH CONVERSION")
    print(f"  Source:   {source_dir}/")
    print(f"  Output:   {output_dir}/")
    print(f"  Files:    {total} total ({len(md_files)} .md, {len(html_files)} .html)")
    print(f"{'='*50}")

    # Create output directory
    ensure_output_dir(os.path.join(output_dir, "placeholder.docx"))

    success_count = 0
    fail_count = 0

    # Convert markdown files
    for filename in md_files:
        try:
            input_path = os.path.join(source_dir, filename)
            base_name = os.path.splitext(filename)[0]
            output_path = os.path.join(output_dir, f"{base_name}.docx")
            convert_to_docx(input_path, output_path)
            success_count += 1
        except Exception as e:
            print(f"  [FAILED] {filename} — {e}")
            fail_count += 1

    # Convert HTML files
    for filename in html_files:
        try:
            input_path = os.path.join(source_dir, filename)
            base_name = os.path.splitext(filename)[0]
            output_path = os.path.join(output_dir, f"{base_name}.docx")
            convert_to_docx(input_path, output_path)
            success_count += 1
        except Exception as e:
            print(f"  [FAILED] {filename} — {e}")
            fail_count += 1

    print(f"\n{'='*50}")
    print(f"  BATCH COMPLETE: {success_count} succeeded, {fail_count} failed")
    print(f"{'='*50}")


# --- CLI Entry ---
def build_parser():
    """Build argument parser with subcommands."""
    parser = argparse.ArgumentParser(
        description="Convert Markdown or HTML wiki pages to Word (.docx) format.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Convert a single wiki page (auto-resolves wiki/log.md → export/log.docx)
  python md_or_html_to_word.py log

  # Convert with explicit paths
  python md_or_html_to_word.py wiki/log.md export/log.docx

  # Batch convert all wiki pages
  python md_or_html_to_word.py --batch

  # Batch with custom directories
  python md_or_html_to_word.py --batch --source raw --output docs
        """,
    )

    parser.add_argument(
        "input",
        nargs="?",
        help="Input file (e.g., 'log', 'wiki/log.md', 'log.html'). "
             "If no extension, searches wiki/<name>.*",
    )

    parser.add_argument(
        "output",
        nargs="?",
        help="Output .docx file path (default: export/<input_name>.docx)",
    )

    parser.add_argument(
        "--batch",
        action="store_true",
        help="Batch convert all files in the source directory",
    )

    parser.add_argument(
        "--source",
        default=None,
        help=f"Source directory for --batch mode (default: {DEFAULT_SOURCE_DIR})",
    )

    parser.add_argument(
        "--output-dir",
        default=None,
        help=f"Output directory for --batch mode and single file default (default: {DEFAULT_OUTPUT_DIR})",
    )

    return parser


def main():
    parser = build_parser()
    args = parser.parse_args()

    # Batch mode
    if args.batch:
        source_dir = args.source or DEFAULT_SOURCE_DIR
        output_dir = args.output_dir or DEFAULT_OUTPUT_DIR
        batch_convert(source_dir, output_dir)
        return

    # Single file mode
    if args.input is None:
        parser.print_help()
        print("\n[WARNING] Please specify an input file or use --batch")
        sys.exit(1)

    input_path = args.input

    # Determine output path
    if args.output:
        output_path = args.output
    else:
        output_dir = args.output_dir or DEFAULT_OUTPUT_DIR
        # Extract base name (without extension)
        base_name = os.path.splitext(os.path.basename(input_path))[0]
        # If input has directory prefix, strip it for output filename
        base_name = os.path.basename(base_name)
        output_path = os.path.join(output_dir, f"{base_name}.docx")

    # Resolve input if needed (for convenience naming)
    resolved = resolve_input_path(input_path)
    if resolved is None:
        print(f"\n[ERROR] Input file not found: {input_path}")
        print(f"\n  Searched:")
        print(f"    {input_path}")
        print(f"    {DEFAULT_SOURCE_DIR}/{input_path}")
        for ext in [".md", ".markdown", ".html", ".htm"]:
            print(f"    {DEFAULT_SOURCE_DIR}/{input_path}{ext}")
        print(f"\n  Available wiki pages:")
        if os.path.isdir(DEFAULT_SOURCE_DIR):
            for f in sorted(os.listdir(DEFAULT_SOURCE_DIR)):
                if f.endswith((".md", ".markdown", ".html", ".htm")):
                    name = os.path.splitext(f)[0]
                    print(f"    - {name}")
        sys.exit(1)

    # Perform conversion
    try:
        convert_to_docx(resolved, output_path)
        print(f"\n[SUCCESS] Conversion complete!\n")
    except Exception as e:
        print(f"\n[ERROR] Error during conversion:\n   {e}\n")
        sys.exit(1)


if __name__ == "__main__":
    main()